from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "学号姓名_深入学习党的二十届四中全会精神.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, first_line=True, before=0, after=6, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, first_line=False, before=10, after=6, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 14, bold=True, color=(151, 48, 38))
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("深入学习党的二十届四中全会精神")
    set_run_font(title_run, size=22, bold=True, color=(151, 48, 38))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(12)
    subtitle_format.space_after = Pt(36)
    sub_run = subtitle.add_run("思想政治理论课学习作业")
    set_run_font(sub_run, size=14, bold=False, color=(89, 89, 89))

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_table.autofit = False
    labels = ["作业题目", "学号", "姓名"]
    values = ["深入学习党的二十届四中全会精神", "请填写学号", "请填写姓名"]
    widths = [Cm(3.2), Cm(10.2)]
    for row_idx, row in enumerate(info_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = widths[col_idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "single")
                tag.set(qn("w:sz"), "8")
                tag.set(qn("w:space"), "0")
                tag.set(qn("w:color"), "BFBFBF")
                tc_borders.append(tag)
            tc_pr.append(tc_borders)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph(p, first_line=False, after=0, line=1.4)
            r = p.add_run(labels[row_idx] if col_idx == 0 else values[row_idx])
            set_run_font(r, size=12, bold=(col_idx == 0))
            if col_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F7E6E3")
                tc_pr.append(shd)

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(30)
    p_note_run = p_note.add_run("文件命名：学号+姓名（例：222219802122张三）")
    set_run_font(p_note_run, size=10.5, color=(89, 89, 89))

    doc.add_section(WD_SECTION.NEW_PAGE)

    # Body title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("深入学习党的二十届四中全会精神")
    set_run_font(r, size=18, bold=True, color=(151, 48, 38))

    lead = doc.add_paragraph()
    set_paragraph(lead, first_line=False, before=0, after=12, line=1.5)
    add_shading(lead, "FDF2F0")
    lr = lead.add_run(
        "党的二十届四中全会于2025年10月20日至23日在北京举行，审议通过了"
        "《中共中央关于制定国民经济和社会发展第十五个五年规划的建议》。"
        "深入学习全会精神，有助于我们把个人成长放到国家发展大局中思考，"
        "进一步明确新时代青年应有的理想追求和实践方向。"
    )
    set_run_font(lr, size=12, bold=False)

    add_heading(doc, "一、时代背景：在承前启后的关键时期谋划新发展")
    add_body(
        doc,
        "党的二十届四中全会是在我国即将完成“十四五”规划主要目标任务、"
        "进入“十五五”时期的重要节点召开的会议。过去五年，我国面对复杂严峻的国际形势和艰巨繁重的改革发展稳定任务，"
        "坚持稳中求进，推动经济实力、科技实力、综合国力跃上新台阶，中国式现代化迈出新的坚实步伐。"
    )
    add_body(
        doc,
        "同时也要看到，当前世界百年变局加速演进，新一轮科技革命和产业变革深入发展，"
        "外部环境中的不确定性、不稳定性仍然较多。国内发展也面临有效需求不足、关键核心技术攻关、区域协调、"
        "生态治理、民生保障等方面的新课题。因此，全会把“十五五”时期放在基本实现社会主义现代化进程中来谋划，"
        "既总结既有成就，又回应现实挑战，体现了党中央对发展阶段、发展条件和发展任务的深刻把握。"
    )

    add_heading(doc, "二、重大意义：为中国式现代化提供战略指引")
    add_body(
        doc,
        "学习党的二十届四中全会精神，首先要认识到它的战略意义。全会围绕制定“十五五”规划提出建议，"
        "这不仅关系未来五年经济社会发展的方向，也关系到到2035年基本实现社会主义现代化目标能否取得决定性进展。"
        "“十五五”时期具有承前启后的重要地位，既要巩固“十四五”时期形成的发展基础，又要为更长远的现代化建设积蓄力量。"
    )
    add_body(
        doc,
        "其次，全会强调高质量发展、改革创新、人民至上、统筹发展和安全等原则，"
        "为我们理解中国式现代化提供了清晰坐标。中国式现代化不是单纯追求经济总量扩张，"
        "而是经济发展、科技进步、文化繁荣、民生改善、生态文明和国家安全相互支撑的现代化。"
        "这说明国家发展越向前推进，越需要系统思维、底线思维和改革精神。"
    )

    add_heading(doc, "三、基本内容：把握全会精神中的重点问题")
    add_body(
        doc,
        "第一，坚持党的全面领导。全会精神最根本的一点，是坚持党中央集中统一领导，"
        "把党的领导贯穿经济社会发展全过程。只有方向明确、组织有力、步调一致，才能在复杂环境中保持战略定力，"
        "集中力量办好自己的事。"
    )
    add_body(
        doc,
        "第二，坚持高质量发展。全会提出建设现代化产业体系，巩固壮大实体经济根基，"
        "加快高水平科技自立自强，引领发展新质生产力。对青年学生来说，这启示我们不能只满足于掌握书本知识，"
        "还要关注科技创新、产业升级和数字中国建设，提升解决实际问题的能力。"
    )
    add_body(
        doc,
        "第三，坚持以人民为中心。全会强调加强普惠性、基础性、兜底性民生建设，"
        "解决好人民群众急难愁盼问题，推动人的全面发展、全体人民共同富裕迈出坚实步伐。"
        "发展的最终目的不是抽象的数字增长，而是让人民生活更加幸福美好。"
    )
    add_body(
        doc,
        "第四，坚持全面深化改革和高水平开放。改革是破解发展难题的关键一招，开放是中国式现代化的重要动力。"
        "无论是建设全国统一大市场，还是推动贸易创新发展、高质量共建“一带一路”，都体现了在更高水平上配置资源、"
        "激发活力、拓展空间的要求。"
    )
    add_body(
        doc,
        "第五，坚持统筹发展和安全。全会对推进国家安全体系和能力现代化、建设更高水平平安中国作出部署。"
        "在现代化进程中，发展是安全的基础，安全是发展的保障。面对网络安全、科技安全、粮食能源安全、社会治理等问题，"
        "必须增强忧患意识和责任意识。"
    )

    add_heading(doc, "四、实践要求：把学习成果转化为实际行动")
    add_body(
        doc,
        "作为新时代青年，学习全会精神不能停留在口号和概念上，而应落实到日常学习、能力提升和责任担当中。"
        "首先，要坚定理想信念，增强对中国特色社会主义道路、理论、制度、文化的自信，"
        "把个人理想同国家前途、民族命运联系起来。"
    )
    add_body(
        doc,
        "其次，要练就过硬本领。面对科技革命和产业变革，青年学生应珍惜学习时间，夯实专业基础，"
        "提高信息素养、创新意识和实践能力，努力成为能适应未来发展需要的高素质人才。"
        "只有把知识学扎实、把能力练过硬，才能在国家需要的时候贡献自己的力量。"
    )
    add_body(
        doc,
        "再次，要树立服务人民、服务社会的意识。无论将来从事什么岗位，都应把个人奋斗放在社会需要之中，"
        "关注基层实际，关心民生问题，主动参加社会实践、志愿服务和集体活动，在实践中理解国情、增长才干、磨炼品格。"
    )
    add_body(
        doc,
        "最后，要保持求真务实的作风。学习贯彻全会精神，关键在于知行合一。"
        "我们应从按时完成学习任务、遵守纪律、认真对待每一次实践训练做起，把爱国情、强国志、报国行统一起来，"
        "努力在推进中国式现代化的新征程中展现青年担当。"
    )

    add_heading(doc, "结语")
    add_body(
        doc,
        "党的二十届四中全会为“十五五”时期发展擘画了蓝图，也为青年成长指明了方向。"
        "通过学习，我更加认识到，国家发展不是遥远的宏大叙事，而与每个人的学习、工作和生活紧密相关。"
        "新时代青年应以更加积极的态度学习理论、锤炼本领、投身实践，在强国建设、民族复兴的伟大进程中贡献青春力量。"
    )

    add_heading(doc, "参考资料", level=2)
    refs = [
        "1. 新华社：《中国共产党第二十届中央委员会第四次全体会议公报》，2025年10月23日。",
        "2. 中国政府网：《中国共产党第二十届中央委员会第四次全体会议公报》，2025年10月23日。",
        "3. 新华网：《中宣部组织召开学习宣传贯彻党的二十届四中全会精神电视电话会议》，2025年10月23日。",
    ]
    for ref in refs:
        rp = doc.add_paragraph()
        set_paragraph(rp, first_line=False, after=3, line=1.25)
        rr = rp.add_run(ref)
        set_run_font(rr, size=10.5, color=(89, 89, 89))

    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("深入学习党的二十届四中全会精神")
    set_run_font(fr, size=9, color=(128, 128, 128))

    doc.save(OUTPUT)


if __name__ == "__main__":
    create_doc()
